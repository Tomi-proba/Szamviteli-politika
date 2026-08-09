#!/usr/bin/env python3
"""
Tartalmi ellenőrzés: a legenerált vázlatban tényleg a helyes VAGY-ág maradt
bent, és a behelyettesített válaszok tényleg a sablon "kitöltendő hely"
jelölésével (sárga kiemelés + áthúzás + piros betű) szerepelnek.

A paritás-teszt (parity_check.py) csak azt garantálja, hogy a két motor
UGYANAZT adja; ez a teszt azt, hogy a közös kimenet HELYES is.

Futtatás:
    python3 scripts/dev/content_check.py
"""
import json
import subprocess
import sys
import tempfile
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(BASE / "scripts"))

import docx  # noqa: E402
from docx.table import _Cell  # noqa: E402

import generate_policy  # noqa: E402

MARK = ("YELLOW (7)", True, "C00000")  # a mark_substituted() jelölése


def _is_marked(r):
    color = None
    if r.font.color is not None and r.font.color.type is not None:
        color = str(r.font.color.rgb)
    return (str(r.font.highlight_color), bool(r.font.strike), color) == MARK


def load(answers_path):
    """A legenerált vázlat bekezdés-szövegei, a behelyettesítettként jelölt
    futások, valamint a TÁBLÁZATOK tartalma (cellaszövegek soronként).

    A táblázatokat a doc.paragraphs egyáltalán nem tartalmazza, ezért a
    táblázatba kitöltött válaszokat külön kell összegyűjteni."""
    out = Path(tempfile.mkdtemp()) / "o.docx"
    generate_policy.generate(answers_path, out)
    doc = docx.Document(str(out))
    paras = [p.text for p in doc.paragraphs]
    marked = []
    for p in doc.paragraphs:
        for r in p.runs:
            if _is_marked(r):
                marked.append(r.text)
    tables = []
    for tbl in doc.tables:
        rows = []
        for tr in generate_policy.table_rows(tbl):
            cells = []
            for tc in generate_policy.row_cells(tr):
                cell = _Cell(tc, tbl)
                cells.append(" ".join(p.text for p in cell.paragraphs).strip())
                for p in cell.paragraphs:
                    for r in p.runs:
                        if _is_marked(r):
                            marked.append(r.text)
            rows.append(cells)
        tables.append(rows)
    return paras, marked, tables


def check(answers_path, must_contain, must_not_contain, must_be_marked,
          table_rows_present=(), table_rows_absent=(), table_count=None):
    paras, marked, tables = load(answers_path)
    body = "\n".join(paras)
    flat_rows = [" | ".join(row) for tbl in tables for row in tbl]
    name = Path(answers_path).name
    fails = []
    for s in must_contain:
        if s not in body:
            fails.append(f"HIÁNYZIK a kimenetből: {s!r}")
    for s in must_not_contain:
        if s in body:
            fails.append(f"BENNMARADT, pedig törölni kellett volna: {s!r}")
    for s in must_be_marked:
        if s not in marked:
            fails.append(f"NEM behelyettesítettként jelölt (sárga/áthúzott/piros): {s!r}")
    if table_count is not None and len(tables) != table_count:
        fails.append(f"a táblázatok száma {len(tables)}, várt: {table_count}")
    for s in table_rows_present:
        if not any(s in row for row in flat_rows):
            fails.append(f"HIÁNYZIK a táblázatokból: {s!r}")
    for s in table_rows_absent:
        if any(s in row for row in flat_rows):
            fails.append(f"BENNMARADT a táblázatokban: {s!r}")
    # a VAGY / VÁLASZTANI KELL elválasztók a feloldott csoportokból eltűntek-e
    print(f"  {'HIBA' if fails else 'OK  '}  {name}: {len(paras)} bekezdés, "
          f"{len(marked)} behelyettesített run, {len(tables)} táblázat")
    for f in fails:
        print(f"        - {f}")
    return not fails


def main():
    a1 = BASE / "data" / "answers.example.json"
    a2 = BASE / "data" / "answers.example2.json"
    ans1 = json.loads(a1.read_text(encoding="utf-8"))
    ans2 = json.loads(a2.read_text(encoding="utf-8"))

    ok = True

    # ---- 1. minta: az EREDETI 24 kérdés + az újak "A" ága ----
    ok &= check(
        a1,
        must_contain=[
            # eredeti kérdések behelyettesítései
            ans1["beszamolo_alairoja"],
            ans1["konyveles_felelos_nev_regszam"],
            ans1["konyvvezetesert_felelos_szemely"],
            ans1["konyvelo_program"],
            ans1["jelentos_mertek_hanyad"],
            ans1["kiveteles_nagysag_osszeghatar"],
            ans1["bizonylat_konyveles_osszhang"],
            ans1["bizonylatkezeles_szabalyai"],
            ans1["evkozi_zarasok_feladatai"],
            ans1["kiegeszito_melleklet_sajatossagok"],
            ans1["idegen_nyelv"],
            # eredeti VAGY-ágak
            "megbízott személy mérlegképes könyvelői",      # könyvvezetés-felelős ág
            "Jelenleg a Vállalkozás nem minősül",          # konszolidáció: nem_erintett
            "Interneten nem tesszük közzé",                 # honlap: nem
            # új kérdések ágai
            "a Vállalkozás az üzleti évéről a számviteli törvény szerinti éves beszámolót készít",
            "„A” változat",
            "összköltség eljárású",
            "csak az 5. számlaosztályban",
            "Vállalkozásunk éves beszámolójának kiegészítő melléklete",
        ],
        must_not_contain=[
            "[Szv 32. kérdés]",                             # bracket-hivatkozás
            "JELÖLT MUNKAPÉLDÁNY",                          # szerkesztői útmutató
            "„B” változat",
            "forgalmi költség eljárású",
            "egyszerűsített éves beszámolójának kiegészítő melléklete",
            "a Vállalkozás az üzleti évéről a számviteli törvény szerinti mikrogazdálkodói",
            "elsődlegesen a 6-7. számlaosztályokban",
        ],
        must_be_marked=[
            ans1["beszamolo_alairoja"],
            ans1["konyvelo_program"],
            ans1["idegen_nyelv"],
            ans1["evkozi_zarasok_feladatai"],
            # táblázatba behelyettesített válaszok is jelölve vannak
            ans1["szv1"],
            ans1["cegjegyzekszam"],
            ans1["szv52"],
            ans1["szv75"],
        ],
        # a bizonylat-aláíró ág marad -> mind a 6 táblázat megvan
        table_count=6,
        table_rows_present=[
            # fejléc-táblázat (cégadatok)
            ans1["szv1"], ans1["szv2"], ans1["szv6"], ans1["cegjegyzekszam"], ans1["szv5"],
            # aláíró-táblázat: a sablon 1 üres sorából 2 kitöltött sor lett
            f'Kimenő számla | {ans1["szv52"]}',
            f'Egyéb kimenő számviteli bizonylat | {ans1["szv53"]}',
            # értékvesztési mértékek: a lefedett kategóriák felülírva...
            f'I. | Immateriális javak | {ans1["szv75"]}',
            f'III. | Befektetett pénzügyi eszközök | {ans1["szv76"]}',
            f'IV. | Készletek | {ans1["szv84"]}',
            f'V. | Követelések | {ans1["szv87"]}',
            # ...a kérdéssel nem lefedett kategóriák viszont NEM
            "II. | Tárgyi eszközök | 20 %",
            "VI. | Értékpapírok | 20 %",
            "VIII. | Részesedések | 20 %",
            # a szándékosan érintetlenül hagyott sablon-táblázatok default szövege
            "Készpénzforgalom bizonylatának feldolgozása",
            "Főkönyvi kivonat készítése",
        ],
    )

    # ---- 2. minta: minden ellentétes ág ----
    ok &= check(
        a2,
        must_contain=[
            ans2["beszamolo_alairoja"],
            ans2["konyvvizsgalo_adatai"],
            ans2["ertekelesi_delegalas"],
            ans2["leltar_delegalas_reszletek"],
            ans2["konszolidacio_reszletszabalyok"],
            ans2["beszamolo_honlap_cim"],
            ans2["celtartalek_jelentos_hatar"],
            ans2["bizomanyi_dij_hatar"],
            ans2["cegertek_leiras_ev"],
            "„B” változat",
            "forgalmi költség eljárású",
            "egyszerűsített éves beszámolójának kiegészítő melléklete",
            "elsődlegesen a 6-7. számlaosztályokban",
            "a pénzkezelési szabályzatban megjelölt személyek végzik",
        ],
        must_not_contain=[
            "JELÖLT MUNKAPÉLDÁNY",
            "„A” változat",
            "összköltség eljárású",
            "Vállalkozásunk éves beszámolójának kiegészítő melléklete",
            "Interneten nem tesszük közzé",
            "csak az 5. számlaosztályban, költségnemenként számolja el",
        ],
        must_be_marked=[
            ans2["beszamolo_alairoja"],
            ans2["konyvvizsgalo_adatai"],
            ans2["beszamolo_honlap_cim"],
            ans2["celtartalek_elteres_szazalek"],
            ans2["szv1"],
            ans2["cegjegyzekszam"],
        ],
        # a "képviselő aláírása" ág marad -> az aláíró-táblázatot töröltük,
        # így 6 helyett csak 5 táblázat maradt a dokumentumban
        table_count=5,
        table_rows_present=[
            ans2["szv1"], ans2["szv2"], ans2["szv6"], ans2["cegjegyzekszam"], ans2["szv5"],
            # üres %-válaszoknál MINDEN kategória megtartja a 20%-os alapértéket
            "I. | Immateriális javak | 20 %",
            "III. | Befektetett pénzügyi eszközök | 20 %",
            "IV. | Készletek | 20 %",
            "V. | Követelések | 20 %",
        ],
        table_rows_absent=[
            # az aláíró-táblázat fejléce sem maradhatott a dokumentumban
            "Feljogosított aláíró",
            "Kimenő számla",
        ],
    )

    print("TARTALMI ELLENŐRZÉS: RENDBEN" if ok else "TARTALMI ELLENŐRZÉS: HIBA")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
