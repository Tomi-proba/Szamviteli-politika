#!/usr/bin/env python3
"""
Python <-> JS motor-paritás ellenőrzés.

A generate_policy.py kimenetét és a webapp-client (zip.js + docx.js +
resolve.js) böngészőben futtatott kimenetét hasonlítja össze RUN-szinten:
minden bekezdés minden futásának szövege, kiemelése, áthúzása és betűszíne
meg kell egyezzen. Ez a rendszer alapvető helyességi garanciája - a két
motornak azonos kimenetet kell adnia ugyanarra a válaszkészletre.

Futtatás:
    python3 scripts/dev/parity_check.py data/answers.example.json [...]
"""
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(BASE / "scripts"))

import docx  # noqa: E402
from docx.table import _Cell  # noqa: E402

import generate_policy  # noqa: E402

CLIENT = BASE / "webapp-client"

RUNNER_JS = r"""
const { chromium } = require('playwright');
const fs = require('fs');
const path = require('path');
(async () => {
  const [engineDir, tmplPath, answersPath, outPath] = process.argv.slice(2);
  const engine = ['zip.js', 'docx.js', 'resolve.js']
    .map(n => fs.readFileSync(path.join(engineDir, n), 'utf8')).join('\n');
  const tmplB64 = fs.readFileSync(tmplPath).toString('base64');
  const answers = JSON.parse(fs.readFileSync(answersPath, 'utf8'));
  const browser = await chromium.launch();
  const page = await browser.newPage();
  await page.setContent('<!doctype html><meta charset="utf-8"><body></body>');
  await page.addScriptTag({ content: engine });
  const arr = await page.evaluate(async ({ tmplB64, answers }) => {
    const out = await generateFromAnswers(tmplB64, answers);
    return Array.from(new Uint8Array(out));
  }, { tmplB64, answers });
  fs.writeFileSync(outPath, Buffer.from(arr));
  await browser.close();
})();
"""


def js_generate(answers_path, out_path):
    with tempfile.NamedTemporaryFile("w", suffix=".js", delete=False) as f:
        f.write(RUNNER_JS)
        runner = f.name
    env = dict(os.environ)
    env.setdefault("NODE_PATH", "/opt/node22/lib/node_modules")
    subprocess.run(
        ["node", runner, str(CLIENT), str(BASE / "sablon" / "merged.docx"),
         str(answers_path), str(out_path)],
        check=True, env=env,
    )
    Path(runner).unlink()


def _run_props(r):
    color = None
    if r.font.color is not None and r.font.color.type is not None:
        color = str(r.font.color.rgb)
    return (r.text, str(r.font.highlight_color), str(r.font.strike), color)


def run_fingerprint(path):
    """Minden futás (szöveg, kiemelés, áthúzás, szín) négyese - ez a
    paritás-összehasonlítás alapja.

    A törzs bekezdésein TÚL a táblázatok celláit is bejárja: a táblázatokat
    a python-docx doc.paragraphs egyáltalán nem tartalmazza, így a
    táblázat-kitöltő resolve_* függvények eltérései e nélkül némán
    átcsúsznának a paritás-ellenőrzésen. A táblázat/sor/cella/bekezdés
    indexek is bekerülnek a lenyomatba, ezért egy törölt táblázat, egy
    duplikált sor vagy egy elcsúszott cella is azonnal eltérésként jelenik
    meg."""
    doc = docx.Document(str(path))
    out = []
    for pi, p in enumerate(doc.paragraphs):
        for ri, r in enumerate(p.runs):
            out.append(("p", pi, ri) + _run_props(r))
    for ti, tbl in enumerate(doc.tables):
        for ri, tr in enumerate(generate_policy.table_rows(tbl)):
            for ci, tc in enumerate(generate_policy.row_cells(tr)):
                for pi, p in enumerate(_Cell(tc, tbl).paragraphs):
                    for xi, r in enumerate(p.runs):
                        out.append(("t", ti, ri, ci, pi, xi) + _run_props(r))
    return out


def compare(answers_path):
    tmp = Path(tempfile.mkdtemp())
    py_out, js_out = tmp / "py.docx", tmp / "js.docx"
    generate_policy.generate(answers_path, py_out)
    js_generate(answers_path, js_out)
    a, b = run_fingerprint(py_out), run_fingerprint(js_out)
    if a == b:
        d = docx.Document(str(py_out))
        print(f"  OK  {Path(answers_path).name}: {len(a)} run egyezik "
              f"({len(d.paragraphs)} bekezdés, {len(d.tables)} táblázat)")
        return True
    print(f"  ELTÉRÉS {Path(answers_path).name}: py={len(a)} js={len(b)} run")
    for i in range(max(len(a), len(b))):
        x = a[i] if i < len(a) else None
        y = b[i] if i < len(b) else None
        if x != y:
            print(f"    első eltérés #{i}:\n      py={x}\n      js={y}")
            break
    return False


if __name__ == "__main__":
    paths = sys.argv[1:] or [BASE / "data" / "answers.example.json",
                             BASE / "data" / "answers.example2.json"]
    ok = all(compare(p) for p in paths)
    print("PARITÁS: RENDBEN" if ok else "PARITÁS: HIBA")
    sys.exit(0 if ok else 1)
