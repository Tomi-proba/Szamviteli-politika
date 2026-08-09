#!/usr/bin/env python3
"""
Számviteli politika - automatikus vázlatgenerátor.

Beolvassa a JELÖLT sablont (sablon/szamviteli_politika_JELOLT.docx) és egy
kitöltött válaszkészletet (JSON, ld. data/answers.example.json), majd
legenerál egy tisztított Word-vázlatot:
  - kitölti a kérdőívben megválaszolt helyeket,
  - a VAGY-alternatívák közül a válaszok alapján kiválasztja a megfelelőt és
    törli a többit,
  - eltávolítja a [Szv N. kérdés] / [Ért N. kérdés] hivatkozásokat és a
    JELÖLT munkapéldány szerkesztői útmutatóját,
  - minden behelyettesített választ a sablon eredeti "kitöltendő hely"
    jelölésével (sárga kiemelés, áthúzás, piros betűszín) hagy, hogy a
    felülvizsgáló kolléga elsőre lássa, mi került be automatikusan; a
    bekezdés egyéb (nem behelyettesített) részein megszünteti a kék
    "még döntendő" betűszínt.

FONTOS: ez egy VÁZLATOT készít. A sablon számos, a kérdőívben NEM szereplő
döntési pontot (VAGY-alternatívát) is tartalmaz (pl. eszközértékelési
módszerek a VIII. fejezetben) - ezeket a script szándékosan nem nyúlja
hozzá, mert nincs hozzájuk ügyféladat. A munkatársnak ezeket továbbra is
kézzel kell átnéznie és eldöntenie, a dokumentum kiküldése előtt.
"""
import copy
import json
import re
import sys
import unicodedata
from datetime import date, timedelta
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.table import _Cell

BASE = Path(__file__).resolve().parent.parent
TEMPLATE = BASE / "sablon" / "merged.docx"


# --------------------------------------------------------------------------
# Alacsony szintű run/bekezdés segédfüggvények
# --------------------------------------------------------------------------

def blank_groups(paragraph):
    """A bekezdés futásait (run) logikai 'kitöltendő hely' csoportokba
    rendezi. Egy csoport egy vagy több szomszédos, áthúzott (strike) futásból
    áll, köztük legfeljebb üres futásokkal - ezeket a Word run-daraboló
    viselkedése hozza létre, nem önálló helyek."""
    groups, current = [], []
    for i, r in enumerate(paragraph.runs):
        strike = bool(r.font.strike)
        empty_nonstrike = (not strike) and r.text.strip() == ""
        if strike or empty_nonstrike:
            current.append(i)
        else:
            if current and any(paragraph.runs[j].font.strike for j in current):
                groups.append(current)
            current = []
    if current and any(paragraph.runs[j].font.strike for j in current):
        groups.append(current)
    return groups


def strip_bracket_ref(paragraph):
    """A '[Szv N. kérdés]' / '[Ért N. kérdés]' hivatkozás-futásokat üríti."""
    for r in paragraph.runs:
        if r.font.highlight_color is not None and not r.font.strike and re.search(
            r"\[\s*(Szv|Ért)", r.text
        ):
            r.text = ""


def clear_choice_marks(paragraph, except_runs=()):
    """A 'még döntendő' jelöléseket (sárga kiemelés, áthúzás, kék/piros
    betűszín) törli egy már véglegesített bekezdésben - a except_runs-ban
    megadott run-indexeket kihagyja (pl. amiket épp most jelöltünk meg
    behelyettesítettként)."""
    for i, r in enumerate(paragraph.runs):
        if i in except_runs:
            continue
        r.font.highlight_color = None
        r.font.strike = None
        if r.font.color is not None and r.font.color.type is not None:
            r.font.color.rgb = None


def mark_substituted(run):
    """A sablon eredeti 'kitöltendő hely' jelölését (sárga kiemelés,
    áthúzás, piros betűszín) alkalmazza egy behelyettesített válaszra, hogy
    a felülvizsgáló munkatárs elsőre lássa, mely szövegrészek automatikusan
    kerültek be, és melyeket kell ellenőriznie."""
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.strike = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def fill_blanks(paragraph, values):
    """Sorrendben kitölti a bekezdés kitöltendő helyeit a megadott
    értékekkel; a behelyettesített szöveget a sablon eredeti "kitöltendő"
    jelölésével (sárga kiemelés, áthúzás, piros betű) hagyja, a bekezdés
    egyéb részein pedig megszünteti a "még döntendő" jelölést."""
    groups = blank_groups(paragraph)
    marked_indices = set()
    for grp, val in zip(groups, values):
        first = paragraph.runs[grp[0]]
        first.text = text_value(val)
        mark_substituted(first)
        marked_indices.add(grp[0])
        for j in grp[1:]:
            paragraph.runs[j].text = ""
    strip_bracket_ref(paragraph)
    clear_choice_marks(paragraph, except_runs=marked_indices)


def keep(paragraph):
    """Megtartott, de nem kitöltendő bekezdés véglegesítése (pl. egy
    VAGY-alternatíva statikus szövege)."""
    strip_bracket_ref(paragraph)
    clear_choice_marks(paragraph)


def delete(paragraph):
    p = paragraph._p
    p.getparent().remove(p)


def flag_before(paragraph, note):
    """Jól látható, félkövér piros felülvizsgálati jelölést szúr be az adott
    bekezdés elé - olyan döntési pontokhoz, amikhez nincs kérdőív-adat."""
    new_p = paragraph.insert_paragraph_before()
    run = new_p.add_run(f"⚠ ELLENŐRIZENDŐ (nincs kérdőív-adat): {note}")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(0xC0, 0x00, 0x00)


# --------------------------------------------------------------------------
# Táblázat-szintű segédfüggvények
#
# A motor sokáig KIZÁRÓLAG a törzs közvetlen bekezdéseivel dolgozott
# (doc.paragraphs), a táblázatok tartalmát sem a Python, sem a JS oldal nem
# látta. Az alábbi néhány primitív nyitja meg a táblázatcellákat - a
# webapp-client/docx.js getTables()/getRows()/getCells()/cellParagraph()
# függvényeinek PONTOS párjai, ezért a bejárási szabályuknak bitre egyeznie
# kell a két motorban.
#
# FONTOS: szándékosan a NYERS w:tr / w:tc gyerekeket járjuk be, és NEM a
# python-docx `table.rows[i].cells[j]` rácsát. A python-docx a vízszintesen
# összevont (gridSpan) cellákat a rácsban többször is felsorolja - pl. a
# sablon 5. táblázatának fejlécsora fizikailag 3 db w:tc, a rácsban viszont
# 4 cella -, a JS oldal viszont csak a nyers XML-gyerekeket látja. Ha a két
# motor máshogy indexelné a cellákat, ugyanaz a resolve_* MÁS cellába írna.
# --------------------------------------------------------------------------

def table_rows(table):
    """A táblázat közvetlen w:tr gyerekei."""
    return table._tbl.findall(qn("w:tr"))


def row_cells(tr):
    """Egy sor közvetlen w:tc gyerekei."""
    return tr.findall(qn("w:tc"))


def cell_paragraph(table, row_idx, cell_idx, para_idx=0):
    """Egy táblázatcella közvetlen w:p gyerekei közül a para_idx-edik.

    A visszaadott objektum közönséges python-docx Paragraph, tehát a
    blank_groups() / clear_choice_marks() / keep() / mark_substituted()
    primitívek változtatás nélkül működnek rajta."""
    tc = row_cells(table_rows(table)[row_idx])[cell_idx]
    return _Cell(tc, table).paragraphs[para_idx]


def fill_cell(paragraph, value):
    """Táblázatcella kitöltése a megadott értékkel.

    A törzsbekezdésekkel ELLENTÉTBEN a sablon táblázataiban egyetlen
    áthúzott "kitöltendő hely" futás sincs (ld. README), ezért itt nem
    építhetünk a blank_groups()-ra: a cella első futásának szövegét írjuk
    felül, a többit kiürítjük, teljesen üres cellába pedig új futást hozunk
    létre. A behelyettesített szöveget ugyanazzal a jelöléssel hagyjuk
    (sárga kiemelés, áthúzás, piros betű), mint a törzsben - a
    felülvizsgáló kolléga így a táblázatokban is elsőre látja, mi került be
    automatikusan."""
    text = text_value(value)
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        mark_substituted(runs[0])
        for r in runs[1:]:
            r.text = ""
        clear_choice_marks(paragraph, except_runs={0})
    else:
        mark_substituted(paragraph.add_run(text))


def delete_table(table):
    """Egy teljes táblázat törlése - akkor kell, ha a hozzá tartozó
    VAGY-ágat is töröljük, különben árván maradna a dokumentumban."""
    tbl = table._tbl
    tbl.getparent().remove(tbl)


def clone_row(table, row_idx):
    """A megadott sor másolatát szúrja be közvetlenül az eredeti után.

    A sablon több táblázata egyetlen üres adatsorral érkezik, miközben a
    kérdőívnek több, egymástól elkülönülő válasza van rá (pl. a kimenő
    számla és az egyéb kimenő bizonylat aláírója)."""
    tr = table_rows(table)[row_idx]
    new_tr = copy.deepcopy(tr)
    tr.addnext(new_tr)
    return new_tr


# --------------------------------------------------------------------------
# Segéd: hónapnév -> szám, illusztratív dátum a mérlegkészítési blankokhoz
# --------------------------------------------------------------------------

_MONTHS = {
    "január": 1, "február": 2, "március": 3, "április": 4, "május": 5,
    "június": 6, "július": 7, "augusztus": 8, "szeptember": 9,
    "október": 10, "november": 11, "december": 12,
}


def month_num(value):
    s = str(value).strip().lower()
    if s in _MONTHS:
        return _MONTHS[s]
    m = re.search(r"\d+", s)
    return int(m.group()) if m else 12


def illustrative_date(fordulonap_honap, fordulonap_nap, offset_days):
    """A fordulónap + offset_days illusztratív hónap/nap párja a
    mérlegkészítési bekezdésekhez (nem hivatalos határidő, csak a sablon
    zárójeles példájának kitöltése)."""
    try:
        base = date(2024, month_num(fordulonap_honap), int(fordulonap_nap))
        target = base + timedelta(days=int(offset_days))
        return target.month, target.day
    except Exception:
        return "…", "…"


def text_value(value):
    """Egy válasz szöveggé alakítása behelyettesítés előtt.

    FIGYELEM: a docx.js `textValue()` függvényének PONTOS párja kell legyen.
    A nyers Python `str()` és a JS `String()` NEM ekvivalens minden típusra,
    és a kettő eltérése bitre eltérő dokumentumot eredményezne:

      - `str(20.0)` -> "20.0", a `String(20.0)` viszont -> "20". A JS-ben a
        20.0 és a 20 UGYANAZ az érték, megkülönböztetni nem lehet őket,
        ezért az egész értékű lebegőpontos számokat mindkét oldalon
        egészként írjuk ki.
      - `str(True)` -> "True", a `String(true)` viszont -> "true".

    Az űrlapokról mindig sztring érkezik; ez a helper a kézzel írt vagy
    programmal generált answers.json-ok szám-/logikai értékeit teszi
    biztonságossá.
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def is_number(value):
    try:
        float(str(value).strip().replace(",", "."))
        return True
    except (TypeError, ValueError):
        return False


def _deaccent(s):
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")


_PENZNEM_EGYEB_PLACEHOLDERS = {"egyeb", "egyeb deviza", "egyeb devizanem", "mas", "mas deviza"}


def classify_penznem(raw):
    """A szabadon gépelhető pénznem-mező (Szv#12) besorolása.

    A mezőt korábban zárt select-ként kérdeztük (forint/euro/usd/egyeb_deviza),
    az ügyfél viszont bármit begépelhet (pl. "Forint", "Euró", "GBP"). Ékezet-
    és kis/nagybetű-független illesztést végzünk a három ismert devizára, minden
    más gépelt szöveget "egyeb"-ként kezelünk.

    FIGYELEM: a resolve.js `classifyPenznem()` PONTOS párja kell legyen.
    """
    t = (raw or "").strip()
    tl = _deaccent(t.lower())
    if tl in ("", "forint", "huf"):
        return "forint"
    if tl in ("euro", "eur"):
        return "euro"
    if tl in ("usd", "dollar", "us dollar", "amerikai dollar"):
        return "usd"
    return "egyeb"


def answer(a, key, default=""):
    """Egy válasz kiolvasása ág-választáshoz.

    A HIÁNYZÓ és az ÜRESEN HAGYOTT mezőt egyaránt "nincs válasz"-ként
    kezeljük, ilyenkor a megadott alapértelmezést adjuk vissza. Ez azért
    fontos, mert az iroda-oldalon minden select alapértelmezett értéke az
    üres "-- válasszon --", tehát a kitöltetlen mező ÜRES SZTRINGKÉNT
    érkezik, nem hiányzó kulcsként.

    FIGYELEM: a resolve.js `answer()` függvényének PONTOS párja kell legyen.
    A JS `a.kulcs || alapertelmezes` alak NEM ekvivalens a Python
    `a.get(kulcs, alapertelmezes)`-szel (az előbbi az üres sztringre is az
    alapértelmezést adja, az utóbbi nem) - korábban pontosan emiatt tudott a
    két motor eltérő dokumentumot előállítani ugyanarra a válaszkészletre.
    """
    value = a.get(key)
    if value is None:
        return default
    if isinstance(value, str) and value.strip() == "":
        return default
    return value


# --------------------------------------------------------------------------
# Fő generálási logika
# --------------------------------------------------------------------------

def remove_jelolt_instructions(p):
    """A JELÖLT munkapéldány szerkesztői útmutatóját (0-3. bekezdés) törli -
    ez sosem kerülhet ki az ügyfélhez. Ezt HÍVJUK ELŐSZÖR, mielőtt bármi
    mást törölnénk/kitöltenénk, különben a paragrafus-indexek elcsúsznak."""
    for i in (0, 1, 2, 3):
        delete(p[i])


def resolve_nyelv(p, a):
    if a.get("van_idegen_nyelvu_beszamolo") == "igen":
        fill_blanks(p[600], [a.get("idegen_nyelv", "")])
        delete(p[597])
        delete(p[599])
    else:
        keep(p[597])
        delete(p[599])
        delete(p[600])


def resolve_penznem(p, a):
    raw = answer(a, "penznem_tipus", "forint")
    tipus = classify_penznem(raw)
    if tipus == "forint":
        keep(p[616])
        for i in (618, 620, 622, 624, 625, 627, 629, 631):
            delete(p[i])
    elif tipus in ("euro", "usd"):
        keep(p[620])
        if tipus == "euro":
            keep(p[622])
            delete(p[624])
            delete(p[625])
        else:
            delete(p[622])
            delete(p[624])
            keep(p[625])
        fill_blanks(p[627], [a.get("penznem_atteres_datum", "")])
        delete(p[616])
        delete(p[618])
        delete(p[629])
        delete(p[631])
    else:  # egyeb deviza - szabadon gépelt devizanév
        nev = (a.get("penznem_egyeb_neve") or "").strip()
        if not nev:
            raw_norm = _deaccent(raw.strip().lower()).replace("_", " ")
            if raw.strip() and raw_norm not in _PENZNEM_EGYEB_PLACEHOLDERS:
                nev = raw.strip()
        fill_blanks(p[631], [nev])
        for i in (616, 618, 620, 622, 624, 625, 627, 629):
            delete(p[i])
    keep(p[633])


def resolve_merlegkeszites(p, a):
    nap = a.get("merlegkeszites_nap", "")
    honap = a.get("fordulonap_honap", "")
    fnap = a.get("fordulonap_nap", "")
    ill_honap, ill_nap = illustrative_date(honap, fnap, nap if is_number(nap) else 0)
    if a.get("merlegkeszites_van_kivetel") != "igen":
        fill_blanks(p[662], [nap, ill_honap, ill_nap])
        delete(p[664])
        delete(p[666])
        delete(p[667])
        delete(p[668])
    else:
        fill_blanks(p[666], [nap, ill_honap, ill_nap])
        fill_blanks(p[667], [a.get("merlegkeszites_kivetelek", "")])
        delete(p[662])
        delete(p[664])
        delete(p[668])


def resolve_fordulonap(p, a):
    # 651 = statikus "december 31" alapértelmezés, 653 = VAGY, 655 = kitöltendő
    # egyedi hónap/nap - a kérdőívből mindig van konkrét válasz, ezért mindig
    # az egyedi (kitöltött) változatot tartjuk meg.
    fill_blanks(p[655], [a.get("fordulonap_honap", ""), a.get("fordulonap_nap", "")])
    delete(p[651])
    delete(p[653])


def resolve_alairo(p, a):
    fill_blanks(p[369], [a.get("beszamolo_alairoja", "")])


def resolve_konyvvizsgalat(p, a):
    tipus = answer(a, "konyvvizsgalat_tipus", "nincs_kotelezettseg")
    adatai = a.get("konyvvizsgalo_adatai", "")
    branches = {"kotelezo_altalanos": 405, "kotelezo_koztartozas_miatt": 409, "onkentes": 413}
    if tipus == "nincs_kotelezettseg":
        keep(p[401])
        for i in (405, 409, 413):
            delete(p[i])
    else:
        chosen = branches[tipus]
        fill_blanks(p[chosen], [adatai])
        for i in (401, 405, 409, 413):
            if i != chosen:
                delete(p[i])
    for i in (403, 407, 411):
        delete(p[i])

    fenn = answer(a, "fenntarthatosagi_jelentes", "nem_koteles")
    delete(p[415])
    if fenn == "nem_koteles":
        delete(p[417])
        delete(p[419])
        delete(p[421])
    elif fenn == "igen_ugyanaz_a_konyvvizsgalo":
        keep(p[417])
        delete(p[419])
        delete(p[421])
    else:
        delete(p[417])
        delete(p[419])
        keep(p[421])


def resolve_ertekelesi_felelos(p, a):
    if answer(a, "ertekeles_felelos_tipus") == "egyedi_delegalas":
        fill_blanks(p[338], [a.get("ertekelesi_delegalas", "")])
        delete(p[334])
        delete(p[336])
    else:
        keep(p[334])
        delete(p[336])
        delete(p[338])


def resolve_konyveles_felelos(p, a):
    tipus = answer(a, "konyveles_felelos_tipus", "mentesseg_20mft_arbevetel_alatt")
    nev = a.get("konyveles_felelos_nev_regszam", "")
    if tipus == "merlegkepes_konyvelo":
        fill_blanks(p[313], [nev])
        keep(p[312])
        delete(p[317])
        delete(p[318])
        delete(p[322])
    elif tipus == "oklev_konyvvizsgalo":
        fill_blanks(p[318], [nev])
        keep(p[317])
        delete(p[312])
        delete(p[313])
        delete(p[322])
    else:
        keep(p[322])
        delete(p[312])
        delete(p[313])
        delete(p[317])
        delete(p[318])
    delete(p[315])
    delete(p[320])
    keep(p[324])
    fill_blanks(p[964], [a.get("konyvvezetesert_felelos_szemely", "")])
    fill_blanks(p[962], [a.get("konyvelo_program", "")])


def resolve_nyilvanossagra_hozatal(p, a):
    keep(p[441])
    keep(p[443])
    if a.get("beszamolo_honlapon") == "igen":
        fill_blanks(p[445], [a.get("beszamolo_honlap_cim", "")])
        delete(p[449])
    else:
        keep(p[449])
        delete(p[445])
    delete(p[447])


def resolve_konszolidacio(p, a):
    delete(p[844])  # "VÁLASZTANI KELL:" - szerkesztői utasítás
    tipus = answer(a, "konszolidacio_tipus", "nem_erintett")
    reszlet = a.get("konszolidacio_reszletszabalyok", "")
    branch_map = {
        "leany_mentesul_bevonas_alol": ([850], 851),
        "leany_bevonva": ([855], 856),
        "anya_mentesul_folerendelt_anya_miatt": ([860], 861),
        "anya_mentesul_nagysagrend_alapjan": ([865, 868], 869),
        "anya_keszit_konszolidaltat": ([873, 874], 875),
    }
    all_idx = [846, 850, 851, 855, 856, 860, 861, 865, 868, 869, 873, 874, 875]
    if tipus == "nem_erintett":
        keep(p[846])
        for i in all_idx:
            if i != 846:
                delete(p[i])
    else:
        keep_static, fill_idx = branch_map[tipus]
        for i in keep_static:
            keep(p[i])
        fill_blanks(p[fill_idx], [reszlet])
        for i in all_idx:
            if i != fill_idx and i not in keep_static:
                delete(p[i])
    for i in (848, 853, 858, 863, 871):
        delete(p[i])


def resolve_leltar(p, a):
    tipus = answer(a, "leltar_felelos_tipus", "vezetes_vagy_szv_rendert_felelos")
    if tipus == "vezetes_vagy_szv_rendert_felelos":
        keep(p[342])
        delete(p[346])
        delete(p[350])
    elif tipus == "eszkoz_forras_felelosok":
        keep(p[346])
        delete(p[342])
        delete(p[350])
    else:
        fill_blanks(p[350], [a.get("leltar_delegalas_reszletek", "")])
        delete(p[342])
        delete(p[346])
    delete(p[344])
    delete(p[348])


def resolve_simple_fills(p, a):
    # 831 = "VÁLASZTANI KELL:" szerkesztői utasítás, 833 = kitöltendő változat,
    # 835 = VAGY, 837 = fix %-os alternatíva - a kérdőív mindig konkrét
    # összeghatárt kér, ezért mindig a kitöltött változatot tartjuk meg.
    fill_blanks(p[833], [a.get("jelentos_osszeg_bemutatas", "")])
    delete(p[831])
    delete(p[835])
    delete(p[837])

    # 916 = kitöltendő változat, 918 = VAGY, 920 = fix %-os alternatíva.
    fill_blanks(p[916], [a.get("jelentos_osszeg_ertelmezes", "")])
    delete(p[918])
    delete(p[920])

    fill_blanks(p[908], [a.get("jelentos_mertek_hanyad", "")])
    fill_blanks(p[937], [a.get("kiveteles_nagysag_osszeghatar", "")])
    fill_blanks(p[939], [a.get("kiveteles_nagysag_sajattoke_szazalek", "")])
    fill_blanks(p[941], [
        a.get("kiveteles_nagysag_bevetel_szazalek", ""),
        a.get("kiveteles_nagysag_koltseg_szazalek", ""),
    ])

    # 1054/1058 = két fix eljárási alternatíva, 1062 = kitöltendő, részletes
    # változat - a kérdőív ehhez kér konkrét választ, ezért ezt tartjuk meg.
    fill_blanks(p[1062], [a.get("bizonylat_konyveles_osszhang", "")])
    delete(p[1054])
    delete(p[1056])
    delete(p[1058])
    delete(p[1060])

    fill_blanks(p[1015], [a.get("bizonylatkezeles_szabalyai", "")])
    # 1153-nak KÉT kitöltendő helye van: egy listafolytatás ("...archiválás,
    # ....)") és az alábbiakban részletezendő tényleges válasz - az elsőt egy
    # semleges "stb."-vel zárjuk, mert a kérdőívnek csak a másodikra van
    # konkrét adata.
    fill_blanks(p[1153], ["stb.", a.get("evkozi_zarasok_feladatai", "")])
    fill_blanks(p[823], [a.get("kiegeszito_melleklet_sajatossagok", "")])


def resolve_alapitas_atszervezes(p, a):
    if a.get("alapitas_atszervezes_aktivalja") == "igen" and is_number(
        a.get("alapitas_atszervezes_leiras_ev")
    ):
        fill_blanks(p[1810], [a.get("alapitas_atszervezes_leiras_ev", "")])
        delete(p[1805])
        delete(p[1807])
        delete(p[1809])
    else:
        flag_before(
            p[1805],
            "az alapítás-átszervezés aktiválásáról/leírási idejéről szóló válasz "
            "nem egyértelmű vagy 'nem aktiválunk' - kérjük kézzel pontosítani, "
            "ill. ellenőrizni, hogy a kapcsolódó bekerülésiérték-szabály (VIII. "
            "fejezet) is összhangban van-e.",
        )


def resolve_cegertek(p, a):
    if a.get("van_cegertek") == "igen" and is_number(a.get("cegertek_leiras_ev")):
        fill_blanks(p[1815], [a.get("cegertek_leiras_ev", "")])
        delete(p[1812])
        delete(p[1814])
        delete(p[1817])
        delete(p[1818])
    else:
        flag_before(
            p[1812],
            "nincs üzleti/cégérték, vagy a leírási idő nem egyértelmű - "
            "kérjük kézzel törölni/pontosítani a nem releváns bekezdéseket.",
        )

    if a.get("van_cegertek") == "igen" and a.get("van_negativ_cegertek") == "igen" and is_number(
        a.get("negativ_cegertek_leiras_ev")
    ):
        fill_blanks(p[1699], [a.get("negativ_cegertek_leiras_ev", "")])
    else:
        flag_before(
            p[1699],
            "nincs negatív üzleti/cégérték, vagy a leírási idő nem egyértelmű - "
            "kérjük kézzel törölni/pontosítani.",
        )


def resolve_ertekhelyesbites(p, a):
    specs = [
        ("ertekhelyesbites_immat_van", "ertekhelyesbites_immat_javak", 1274, 1276, 1278),
        ("ertekhelyesbites_targyi_van", "ertekhelyesbites_targyi_eszkoz", 1299, 1301, 1303),
        ("ertekhelyesbites_penzugyi_van", "ertekhelyesbites_penzugyi_eszkoz", 1344, 1346, 1348),
    ]
    for gate_key, text_key, fill_idx, sep_idx, static_idx in specs:
        if a.get(gate_key) == "igen":
            fill_blanks(p[fill_idx], [a.get(text_key, "")])
            delete(p[static_idx])
        else:
            keep(p[static_idx])
            delete(p[fill_idx])
        delete(p[sep_idx])


def resolve_celtartalek(p, a):
    if a.get("celtartalek_kepez") == "igen":
        fill_blanks(p[1608], [""])
        delete(p[1612])
        fill_blanks(p[1614], [a.get("celtartalek_jelentos_hatar", "")])
        fill_blanks(p[1630], [a.get("celtartalek_elteres_szazalek", "")])
    else:
        keep(p[1612])
        delete(p[1608])
        delete(p[1614])
        flag_before(
            p[1630],
            "a vállalkozás nem képez céltartalékot lehetőség alapján - "
            "kérjük ellenőrizni, hogy ez a bekezdés (lényeges eltérés %-a) "
            "továbbra is szükséges-e, vagy törlendő.",
        )
    delete(p[1610])


def resolve_beszamolo_forma(p, a):
    """Szv#13 - a beszámoló választott formája.

    Két, ugyanabból a válaszból levezethető döntés:
      - IV.2. fejezet: 487 = éves, 490 = egyszerűsített éves, 493 =
        mikrogazdálkodói (489/492 a VAGY elválasztók);
      - IV.8. fejezet: a kiegészítő melléklet tartalma két teljes,
        egymást kizáró blokkban van leírva (678-773 = éves beszámoló,
        775-814 = egyszerűsített éves beszámoló), 676 = "VÁLASZTANI KELL:",
        774 = VAGY.

    IFRS szerinti beszámolóhoz a sablonnak nincs ága, a mikrogazdálkodói
    beszámolónak pedig egyáltalán nincs kiegészítő melléklete (398/2012.
    Korm. r.) - ezekben az esetekben nem tippelünk, hanem jelölünk.
    """
    tipus = answer(a, "szv13")
    branches = {
        "eves_beszamolo": 487,
        "egyszerusitett_eves_beszamolo": 490,
        "mikrogazdalkodoi_egyszerusitett_eves_beszamolo": 493,
    }
    if tipus in branches:
        chosen = branches[tipus]
        keep(p[chosen])
        for i in (487, 490, 493):
            if i != chosen:
                delete(p[i])
        delete(p[489])
        delete(p[492])
    else:
        flag_before(
            p[485],
            "a beszámoló választott formája nem egyértelmű (pl. IFRS szerinti "
            "beszámoló, amihez ez a sablon nem tartalmaz alternatívát) - kérjük "
            "kézzel kiválasztani a megfelelő bekezdést, a többit törölni.",
        )

    if tipus == "eves_beszamolo":
        for i in range(678, 774):
            keep(p[i])
        for i in range(774, 815):
            delete(p[i])
    elif tipus == "egyszerusitett_eves_beszamolo":
        for i in range(678, 775):
            delete(p[i])
        for i in range(775, 815):
            keep(p[i])
    else:
        flag_before(
            p[676],
            "a kiegészítő melléklet tartalmához (éves / egyszerűsített éves "
            "beszámoló változat) nincs egyértelmű kérdőív-adat - mikrogazdálkodói "
            "beszámolónak egyáltalán nincs kiegészítő melléklete. Kérjük kézzel "
            "kiválasztani a megfelelő blokkot, a másikat törölni.",
        )
        return
    delete(p[676])


def resolve_merleg_forma(p, a):
    # Szv#14 - 525 = "A" változat, 526 = VAGY, 527 = "B" változat.
    tipus = answer(a, "szv14")
    if tipus == "a_valtozat":
        keep(p[525])
        delete(p[527])
    elif tipus == "b_valtozat":
        keep(p[527])
        delete(p[525])
    else:
        flag_before(p[524], "a mérleg választott formája (A/B változat) hiányzik.")
        return
    delete(p[526])


def resolve_eredmenykimutatas_forma(p, a):
    # Szv#15 - 537 = összköltség, 538 = VAGY, 539 = forgalmi költség eljárású.
    # (531-534 a törvény MINDKÉT eljárást ismertető leírása, ott a "VAGY" a
    # jogszabályi felsorolás része, nem választási pont - nem nyúlunk hozzá.)
    tipus = answer(a, "szv15")
    if tipus == "osszkoltseg_eljarassal":
        keep(p[537])
        delete(p[539])
    elif tipus == "forgalmi_koltseg_eljarassal":
        keep(p[539])
        delete(p[537])
    else:
        flag_before(p[536], "az eredménykimutatás választott formája hiányzik.")
        return
    keep(p[536])
    keep(p[540])
    delete(p[538])


def resolve_letetbehelyezes(p, a):
    """Szv#19 - letétbe helyezési határidő. 427 = általános (ötödik hónap),
    429 = VAGY, 431 = szabályozott piacra bevezetett értékpapír esetén
    (negyedik hónap). A 431 csak egy közbeszúrt feltételes tagmondat: a
    mondat alanya és állítmánya a 427-ben, illetve a 433-ban van, ezért ha a
    szabályozott piaci ág kell, azt nem lehet gépi bekezdés-törléssel
    előállítani - ilyenkor jelölünk."""
    if answer(a, "letetbehelyezes_hatarido_tipus", "altalanos") == "altalanos":
        keep(p[427])
        delete(p[429])
        delete(p[431])
    else:
        flag_before(
            p[427],
            "a vállalkozás értékpapírjait EGT-állam szabályozott piacán "
            "forgalmazzák, ezért a letétbe helyezés határideje a negyedik hónap "
            "utolsó napja - kérjük a bekezdést kézzel összevonni/pontosítani.",
        )


def resolve_penzkezeles_felelos(p, a):
    # Szv#37 - 354 = vezetés által írásban kijelölt személyek, 356 = VAGY,
    # 358 = a pénzkezelési szabályzatban megjelölt személyek.
    if answer(a, "penzkezeles_felelos_tipus", "vezetes_altal_kijelolt") == "penzkezelesi_szabalyzat_szerint":
        keep(p[358])
        delete(p[354])
    else:
        keep(p[354])
        delete(p[358])
    delete(p[356])


def resolve_koltsegelszamolas(p, a):
    # Szv#45 - 968 = csak 5. számlaosztály, 972 = 5. elsődleges + 6-7.
    # másodlagos, 976 = 6-7. elsődleges + 5. másodlagos (970/974 = VAGY).
    branches = {"csak_5": 968, "5_elsodleges": 972, "67_elsodleges": 976}
    chosen = branches.get(answer(a, "koltsegelszamolas_tipus", "csak_5"), 968)
    keep(p[chosen])
    for i in (968, 972, 976):
        if i != chosen:
            delete(p[i])
    delete(p[970])
    delete(p[974])


def resolve_szerzodes_elszamolasi_egyseg(p, a):
    """Szv#83 - a nagy tömegben, sorozatban gyártott termékekre alkalmazzák-e
    a szerződés elszámolási egységére vonatkozó szabályrendszert.

    UGYANAZ a döntés NÉGY helyen szerepel a sablonban (készletek, aktív
    időbeli elhatárolás, passzív időbeli elhatárolás, értékesítés nettó
    árbevétele fejezetek) - mind a négyet ugyanabból a válaszból oldjuk fel,
    különben a dokumentum önmagának mondana ellent."""
    alkalmaz = answer(a, "szerzodes_elszamolasi_egyseg", "nem_alkalmazzuk") == "alkalmazzuk"
    for lead, igen, sep, nem in (
        (1374, 1375, 1376, 1377),
        (1479, 1480, 1482, 1484),
        (1676, 1677, 1679, 1681),
        (1885, 1886, 1887, 1888),
    ):
        keep(p[lead])
        if alkalmaz:
            keep(p[igen])
            delete(p[nem])
        else:
            keep(p[nem])
            delete(p[igen])
        delete(p[sep])


def resolve_ertekpapir_elhatarolas(p, a):
    # Szv#94 - 1431 = elhatároljuk, 1433 = VAGY, 1435 = nem határoljuk el.
    if answer(a, "ertekpapir_kamatkulonbozet_elhatarolas", "nem_hataroljuk_el") == "elhataroljuk":
        keep(p[1431])
        delete(p[1435])
    else:
        keep(p[1435])
        delete(p[1431])
    delete(p[1433])


def resolve_arfolyamveszteseg(p, a):
    """Szv#95 / Ért#5 - a devizakészlettel nem fedezett hiteltartozások
    árfolyamveszteségét halasztott ráfordításként elszámolják-e. 1453 = igen
    (teljes mondat), 1455 = VAGY, 1457 = "nem számoljuk el halasztott
    ráfordításként." - ez utóbbi csak mondatvég, az alany a 1453-ban van,
    ezért a "nem" ágat nem lehet bekezdés-törléssel előállítani."""
    if answer(a, "arfolyamveszteseg_halasztott_raforditas", "nem_szamoljuk_el") == "elszamoljuk":
        keep(p[1453])
        delete(p[1455])
        delete(p[1457])
    else:
        flag_before(
            p[1453],
            "a vállalkozás NEM számolja el halasztott ráfordításként a devizás "
            "hiteltartozások árfolyamveszteségét - a két bekezdést (1453/1457) "
            "kérjük kézzel egy mondattá összevonni.",
        )


def resolve_onkoltsegszamitas(p, a):
    """Szv#106 - a 188. bekezdés kijelenti, hogy a Vállalkozás mentesül az
    önköltségszámítási szabályzat készítése alól. Ha ez nem igaz, a sablonban
    nincs alternatív szöveg, ezért jelölünk - ez a mondat így, ellenőrzés
    nélkül nem mehet ki az ügyfélhez."""
    if answer(a, "onkoltsegszamitasi_szabalyzat", "mentesul") == "mentesul":
        keep(p[188])
    else:
        flag_before(
            p[188],
            "a Vállalkozás NEM mentesül az önköltségszámítási szabályzat "
            "készítése alól (van/kell önköltségszámítási szabályzata) - az alábbi "
            "mentességi bekezdés kézi átírást igényel.",
        )


def resolve_ert4(p, a):
    if a.get("ert4_alkalmazza_elhatarolas") == "igen":
        fill_blanks(p[1494], [a.get("bizomanyi_dij_hatar", "")])
        keep(p[1496])
        delete(p[1500])
    else:
        keep(p[1500])
        delete(p[1494])
        delete(p[1496])
    delete(p[1498])


def resolve_cegadatok(p, a, t):
    """Szv#1/2/5/6 - a dokumentum fejlécében álló cégadat-táblázat (t[0]).

    A táblázat a 7. és 8. bekezdés között áll, hat sorral: CÉGNÉV,
    SZÉKHELY, ADÓSZÁM, CÉGJEGYZÉKSZÁM, "A VÁLLALKOZÁS KÉPVISELETÉRE
    JOGOSULT SZEMÉLY", végül a "a továbbiakban: Vállalkozás" sor (ez utóbbi
    statikus, nem kitöltendő). A bal oszlop a felirat, a jobb oszlop az
    érték - mindegyik értékcella üresen érkezik a sablonban.

    Az első sor mindkét cellája HÁROM bekezdésből áll (üres / szöveg /
    üres), ezért ott a középső (1. indexű) bekezdésbe írunk, hogy az érték
    a "CÉGNÉV" felirattal azonos sorban álljon; a többi értékcellának
    egyetlen bekezdése van."""
    tbl = t[0]
    fill_cell(cell_paragraph(tbl, 0, 1, 1), a.get("szv1", ""))
    fill_cell(cell_paragraph(tbl, 1, 1), a.get("szv2", ""))
    fill_cell(cell_paragraph(tbl, 2, 1), a.get("szv6", ""))
    fill_cell(cell_paragraph(tbl, 3, 1), a.get("cegjegyzekszam", ""))
    fill_cell(cell_paragraph(tbl, 4, 1), a.get("szv5", ""))


def resolve_bizonylat_alairas(p, a, t):
    """Szv#51/52/53 - VII.3.6. A bizonylatok hitelessége.

    1035 = a hitelességet a Vállalkozás képviseletére jogosult személy
    aláírása igazolja, 1037 = VAGY, 1039 = "az egyes ... bizonylat típusok
    esetében az alábbiakban felsorolt, arra feljogosított személyek
    aláírásával igazolható:" - és EHHEZ az ághoz tartozik a t[2] táblázat
    (fejléc: "Bizonylat típus | Feljogosított aláíró", egyetlen üres
    adatsorral).

    A README korábban pontosan ezt a csoportot hozta fel példaként arra,
    miért nem lehetett bekötni: a "képviselő aláírása" ág választásakor a
    1039 bekezdés törlése egy árván maradt, üres táblázatot hagyott volna a
    dokumentumban. A delete_table() primitívvel ez már megoldható.

    A táblázat egyetlen adatsora két, egymástól elkülönülő választ kapna
    (Szv#52 = kimenő számlák, Szv#53 = egyéb kimenő számviteli
    bizonylatok), ezért a sort megduplázzuk. A bizonylattípus-feliratok
    magának az adatbekérőnek a szóhasználatát követik, nem találunk ki új
    kategóriákat."""
    tbl = t[2]
    tipus = answer(a, "bizonylat_hitelesites_tipus", "kepviselo_alairasa")
    if tipus == "feljogositott_szemelyek":
        keep(p[1039])
        delete(p[1035])
        keep(cell_paragraph(tbl, 0, 0))
        keep(cell_paragraph(tbl, 0, 1))
        fill_cell(cell_paragraph(tbl, 1, 0), "Kimenő számla")
        fill_cell(cell_paragraph(tbl, 1, 1), a.get("szv52", ""))
        clone_row(tbl, 1)
        fill_cell(cell_paragraph(tbl, 2, 0), "Egyéb kimenő számviteli bizonylat")
        fill_cell(cell_paragraph(tbl, 2, 1), a.get("szv53", ""))
    else:
        keep(p[1035])
        delete(p[1039])
        delete_table(tbl)
    delete(p[1037])


def resolve_ertekvesztes_mertekek(p, a, t):
    """Szv#75/76/84/87 - a "jelentős" könyv szerinti / piaci érték eltérés
    kategóriánkénti mértékei (t[5]), az 1720. bekezdés felvezetésével
    ("A könyv szerinti érték és a piaci érték közötti különbséget
    jelentősnek tekintjük az alábbiak szerint:").

    A táblázat MINDEN kategóriasora tartalmaz egy 20%-os sablon-alapértéket
    a "Százalék" oszlopban (a 3. oszlop, "Érték (E Ft)", végig üres, és
    nincs is rá kérdőív-adat - nem nyúlunk hozzá).

    Csak azt a NÉGY sort írjuk felül, amelyre az adatbekérőnek célzott
    kérdése van - I. Immateriális javak (Szv#75), III. Befektetett
    pénzügyi eszközök (Szv#76), IV. Készletek (Szv#84), V. Követelések
    (Szv#87) -, és csak akkor, ha az ügyfél tényleg megadott értéket. Üres
    válasznál a sablon 20%-os alapértéke marad, mert az egy érvényes
    alapértelmezés, nem egy kitöltetlen hely. A II. Tárgyi eszközök,
    VI. Értékpapírok, VII. Pénzeszközök és VIII. Részesedések sorokhoz nincs
    kérdés, azokat sosem módosítjuk."""
    tbl = t[5]
    for row_idx, key in ((2, "szv75"), (4, "szv76"), (5, "szv84"), (6, "szv87")):
        value = answer(a, key, "")
        if value != "":
            fill_cell(cell_paragraph(tbl, row_idx, 2), value)


def generate(answers_path, output_path):
    answers = json.loads(Path(answers_path).read_text(encoding="utf-8"))
    answers = {k: v for k, v in answers.items() if not k.startswith("_")}

    doc = docx.Document(str(TEMPLATE))
    # EGYETLEN alkalommal olvassuk ki a bekezdéslistát - minden lenti
    # resolve_* és a törlés is EBBE a listába indexel. A python-docx
    # Paragraph objektumok a mögöttes XML elemre mutatnak, ezért törlés
    # után is stabilak maradnak az itt rögzített indexek; ha bármelyik
    # függvény újra kiolvasná a doc.paragraphs-t, a korábbi törlések miatt
    # minden utána következő index elcsúszna.
    p = doc.paragraphs
    # Ugyanez a táblázatokra: a doc.tables listát is EGYSZER olvassuk ki,
    # mert a resolve_bizonylat_alairas() egy egész táblázatot törölhet, ami
    # utána elcsúsztatná a többi táblázat indexét.
    t = doc.tables

    resolve_nyelv(p, answers)
    resolve_penznem(p, answers)
    resolve_merlegkeszites(p, answers)
    resolve_fordulonap(p, answers)
    resolve_alairo(p, answers)
    resolve_konyvvizsgalat(p, answers)
    resolve_ertekelesi_felelos(p, answers)
    resolve_konyveles_felelos(p, answers)
    resolve_nyilvanossagra_hozatal(p, answers)
    resolve_konszolidacio(p, answers)
    resolve_leltar(p, answers)
    resolve_simple_fills(p, answers)
    resolve_alapitas_atszervezes(p, answers)
    resolve_cegertek(p, answers)
    resolve_ertekhelyesbites(p, answers)
    resolve_celtartalek(p, answers)
    resolve_ert4(p, answers)
    resolve_beszamolo_forma(p, answers)
    resolve_merleg_forma(p, answers)
    resolve_eredmenykimutatas_forma(p, answers)
    resolve_letetbehelyezes(p, answers)
    resolve_penzkezeles_felelos(p, answers)
    resolve_koltsegelszamolas(p, answers)
    resolve_szerzodes_elszamolasi_egyseg(p, answers)
    resolve_ertekpapir_elhatarolas(p, answers)
    resolve_arfolyamveszteseg(p, answers)
    resolve_onkoltsegszamitas(p, answers)
    resolve_cegadatok(p, answers, t)
    resolve_bizonylat_alairas(p, answers, t)
    resolve_ertekvesztes_mertekek(p, answers, t)
    remove_jelolt_instructions(p)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Elkészült: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Használat: generate_policy.py <answers.json> <output.docx>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2])
